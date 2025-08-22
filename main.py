import json
import os
from gigachat import GigaChat
from gigachat.models import Chat, Messages, MessagesRole
from new_prompt import sys_prompt, advantages, clients, politeness
from synergy import synergy
from dotenv import load_dotenv
import random
from docx import Document
from docx.shared import Pt
from datetime import datetime

load_dotenv()
giga_auth = os.getenv("GIGACHAT_AUTH")

total_prompt_tokens = 0
total_completion_tokens = 0
total_tokens = 0

sys_prompt_summary = """
Ты получаешь на вход информацию о трёх продуктах. Для каждого продукта будут даны:
Название
Ключевые преимущества для пользователя (личная ценность/выгоды)
Возможные связи или дополнения с другими продуктами
Задача:
Проанализируй три продукта.
Определи, как их можно связать между собой в единую логическую систему или пакетное предложение.
Построй связку так, чтобы каждый продукт усиливал ценность других.
В ответ верни средней длинны абзац, в котором будешь рассказывать клиенту о том, 
почему для него выгодно использование данных товаров, не сильно повторяйся с информацией, которая была предоставлена тебе.
Требования по стилю написания:
Понятным разговорным языком! Без разметки MarkDown! До 70 слов! Пиши конкретнее, используй числа и пиши конкретную выгоду
Не пиши ничего про общую выгоду и экономию, описывай выгоду только по продуктам отдельно. Не пиши в конце подытоживающее предложение!
Вот тебе пример ОТЛИЧНОГО ответа:
"СберПрайм откроет доступ к выгодным сервисам и упростит ваши службы. 
Накопительный счёт обеспечит доход от свободных средств до 7% годовых. 
Кредитная карта позволит управлять финансами и получать кэшбэк до 5%."
Сохрани стиль и структуру из данного примера
"""

giga = GigaChat(
    credentials=giga_auth,
    scope="GIGACHAT_API_B2B",
    verify_ssl_certs=False
)

print(f"""Выберите номер клиента
        0 - Все клиенты
        1 - Клиент 1 (айтишник)
        2 - Клиент 2 (бабка)
        3 - Клиент 3 (студентик)
        4 - Клиент 4 (темщик)
        5 - Клиент 5 (VIP клиентик)
Введите номер: """)
i = int(input())

# Создаем документ для отчета
doc = Document()
doc.add_heading('Отчет по продуктовым предложениям', 0).style.font.size = Pt(16)
doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc.add_paragraph()

# Генерация описаний продуктов
sdf = []
for prod in advantages:
    messages = [
        Messages(role=MessagesRole.SYSTEM, content=f"{sys_prompt}\n\nИспользуй такой стиль общения:\n{politeness}"),
        Messages(role=MessagesRole.USER, content=f"{prod}\n\n{clients[0]}")
    ]
    payload = Chat(
        messages=messages,
        model='GigaChat-2-Max',
        temperature=0,
        top_p=0.1,
        profanity_check=False
    )
    response = giga.chat(payload)
    sdf.append(response.choices[0].message.content.strip('```') + "\n\n")

with open("output.json", "w", encoding="utf-8") as f:
    json.dump(sdf, f, ensure_ascii=False)

with open("output.json", "r", encoding='utf-8') as file:
    products = json.load(file)

# Функция для обработки одного клиента
def process_client(client_idx, client_name):
    global total_prompt_tokens, total_completion_tokens, total_tokens
    prod = random.sample(range(len(advantages)), 3)  # Выбираем 3 случайных продукта
    relevant_products = []
    product_texts = []

    doc.add_heading(f'Клиент: {client_name}', 1).style.font.size = Pt(14)
    doc.add_paragraph('Выбранные продукты:', style='List Bullet')
    for j in prod:
        product_text = products[j].strip()
        relevant_products.append(product_text + synergy[j])
        product_texts.append(product_text)
        doc.add_paragraph(f'Продукт {j + 1}: {product_text}', style='List Bullet')

    messages = [
        Messages(role=MessagesRole.SYSTEM, content=f"{sys_prompt_summary}"),
        Messages(role=MessagesRole.USER, content=f"Вот список продуктов и их связки с другими продуктами:\n\n{relevant_products}")
    ]
    payload = Chat(
        messages=messages,
        model='GigaChat-2-Max',
        temperature=0,
        top_p=0.1,
        profanity_check=False
    )
    response = giga.chat(payload)
    package_offer = response.choices[0].message.content.strip()

    doc.add_paragraph('Пакетное предложение:')
    doc.add_paragraph(package_offer)

    stage2_prompt_tokens = response.usage.prompt_tokens
    stage2_completion_tokens = response.usage.completion_tokens
    stage2_total_tokens = response.usage.total_tokens
    total_prompt_tokens += stage2_prompt_tokens
    total_completion_tokens += stage2_completion_tokens
    total_tokens += stage2_total_tokens

    doc.add_paragraph('Статистика использования токенов:', style='List Bullet')
    doc.add_paragraph(f'Токены на запрос (prompt_tokens): {stage2_prompt_tokens}', style='List Bullet')
    doc.add_paragraph(f'Токены на ответ (completion_tokens): {stage2_completion_tokens}', style='List Bullet')
    doc.add_paragraph(f'Общее количество токенов (total_tokens): {stage2_total_tokens}', style='List Bullet')
    doc.add_paragraph()

    print(f"\nКлиент: {client_name}")
    print("Выбранные продукты:", prod)
    for text in product_texts:
        print(text)
    print("\nПакетное предложение:")
    print(package_offer)
    print(f"\nТокены на этапе 2:")
    print(f"  Токены на запрос (prompt_tokens): {stage2_prompt_tokens}")
    print(f"  Токены на ответ (completion_tokens): {stage2_completion_tokens}")
    print(f"  Общее количество токенов (total_tokens): {stage2_total_tokens}")

# Обработка клиентов
client_names = ['Айтишник', 'Бабка', 'Студентик', 'Темщик', 'VIP клиентик']
if i == 0:
    for idx in range(5):
        process_client(idx, client_names[idx])
else:
    process_client(i - 1, client_names[i - 1])

# Сохранение общей статистики токенов в документ
doc.add_heading('Общая статистика использования токенов', 1).style.font.size = Pt(14)
doc.add_paragraph(f'Токены на запросы (prompt_tokens): {total_prompt_tokens}', style='List Bullet')
doc.add_paragraph(f'Токены на ответы (completion_tokens): {total_completion_tokens}', style='List Bullet')
doc.add_paragraph(f'Общее количество токенов (total_tokens): {total_tokens}', style='List Bullet')

# Сохранение документа
doc.save('product_offer_report.docx')
print("\nГотово! Отчет сохранен в 'product_offer_report.docx'")